# -*- coding: utf-8 -*-
import os
import glob
from openpyxl import Workbook
import openpyxl
import re
import docx 
from docx.shared import Cm 



T_filePath = glob.glob(os.getcwd() + '/' + 'newfolder' + '/')

heading = [ 
"​1. Creation and God",
"2. Creation and Us",
"3. The Fall",
"4. The Flood",
"5. Abraham’s Covenant",
"6. Joseph",
"7. Moses and the Plagues",
"8. The Ten Commandments",
"9. Presence of God",
"10. The Holiness of God (Leviticus)",
"11. Sold Out to God (Deuteronomy)",
"12. Judges",
"13. God is King (1 Samuel)",
"14. David and Goliath",
"15. Psalm 23",
"16. Confrontation and Confession (Psalm 51)",	
"17. Solomon — the Wise and Foolish (Proverbs)",
"18. Job",
"19. Elijah",
"20. Isaiah and the Holiness of God",
"21. Isaiah 53",
"22. Micah",
"23. Hosea",
"24. Habakkuk, Righteousness and Faith",
"25. The New Covenant (Jeremiah and Ezekiel)",
"26. Lamentations",
"27. The Birth of Jesus",
"28. John the Baptist",
"29. Nicodemus and Rebirth",
"30. The Beatitudes",
"31. The Lord’s Prayer",
"32. Seeking God",
"33. Deity of Christ",
"34. Discipleship",
"35. The Greatest Commandment",
"36. Eschatology",
"37. Holy Spirit",
"38. The Lord’s Supper",
"39. Jesus’ Death and Resurrection",
"40. The Great Commission",
"41. Pentecost",
"42. The Church",
"43. Justification by Faith",
"44. The Grace of Giving",
"45. Christian Joy",
"46. Humility",
"47. Scripture",
"48. Assurance and Perseverance (Hebrews)",
"49. The Tongue (James)",
"50. Suffering and Heaven",
"51. Christian Love",
"52. Revelation",
"53. Revelations" 
]


f = open("52stories.txt", "r")
content = f.read()
split_content = content.split("\n")
translation = " "
count = 0
Chapter = 0
dictnry = {} 
for lines in split_content:
    if lines in heading:
        dictnry[Chapter] = count
        count = 0
        splitLines = lines.split('.')[0]
        print(splitLines)
        Chapter = splitLines
        # print(Chapter)
    else:
        # print(lines)
        rem_numbers = re.sub(r'\d+','',lines)
        rem_punctuation = re.sub(r'[^\w\s]','',rem_numbers)
        # output =  re.sub(r"\b[a-zA-Z]\b", "", rem_punctuation) #remove single letter
        splitLines1 = rem_punctuation.split()
        count += len(splitLines1)
print(dictnry)
    
        

doc = docx.Document()
heading = doc.add_heading("Word Count", level=2).alignment = 1
sections = doc.sections
for section in sections:
    section.top_margin = Cm(1)
    section.bottom_margin = Cm(1)
    section.left_margin = Cm(5)
    section.right_margin = Cm(5)


table = doc.add_table(rows=0, cols=3)
table.style = 'Table Grid'
table.autofit = False 
table.allow_autofit = False 
table.columns[0].width = Cm(1.15) 
table.columns[1].width = Cm(4)
table.columns[2].width = Cm(4)
heading1 = table.add_row().cells
heading1[0].paragraphs[0].add_run('No').bold = True
heading1[1].paragraphs[0].add_run('Book').bold = True
heading1[2].paragraphs[0].add_run('Count').bold = True

count = 1
for k, v in dictnry.items():
    cells = table.add_row().cells
    cells[0].paragraphs[0].add_run(str(count)).font.size = Cm(.34)
    cells[1].paragraphs[0].add_run(str(str(k))).font.size = Cm(.34)
    cells[2].paragraphs[0].add_run(str(str(v))).font.size = Cm(.34)
    count += 1
        
doc.add_page_break()
doc.save('BookCount.docx')
print("saved")