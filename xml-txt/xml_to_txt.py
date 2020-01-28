import re
import os
import glob
import docx
from docx.shared import Cm 
from docx.shared import Pt


files = glob.glob(os.getcwd() + "/files/*")
print files
countw = {}
total_count = 0
# # uniqueWords = []

for fl in files:
    bookname = fl.split('/')[-1]
    rem_usfm = bookname.split('.')[-2]
    # print rem_usfm
    f = open(fl, 'r')
    content = f.read()
    splitcontent = content.split('\n')
    # print splitcontent
    wordcount = []
    for lines in splitcontent:
        rem_tags = re.sub(r'<.*?>', '', lines)
        # rem_numbers = re.sub(r'\d+','',rem_tags)
        # rem_punctuation = re.sub(r'[^\w\s]','',rem_numbers)
        # output =  re.sub(r"\b[a-zA-Z]\b", "", rem_punctuation)
        print rem_tags
#         splitContent = output.split(" ")
#         for words in splitContent:
#             if words == '':
#                 pass
#             else:
#                 wordcount.append(words)
# #                 # uniqueWords.append(words)

#     wordLength = len(wordcount)
#     countw[rem_usfm] = wordLength
#     total_count += wordLength
# print((countw))
# print(total_count)
# # # print(len(set(uniqueWords)))      
            

# doc = docx.Document()
# heading = doc.add_heading("Word Count", level=2).alignment = 1
# sections = doc.sections
# for section in sections:
#     section.top_margin = Cm(1)
#     section.bottom_margin = Cm(1)
#     section.left_margin = Cm(5)
#     section.right_margin = Cm(5)


# table = doc.add_table(rows=0, cols=3)
# table.style = 'Table Grid'
# table.autofit = False 
# table.allow_autofit = False 
# table.columns[0].width = Cm(1.15) 
# table.columns[1].width = Cm(4)
# table.columns[2].width = Cm(4)
# heading1 = table.add_row().cells
# heading1[0].paragraphs[0].add_run('No').bold = True
# heading1[1].paragraphs[0].add_run('Book').bold = True
# heading1[2].paragraphs[0].add_run('Count').bold = True

# count = 1
# for k, v in countw.items():
#     cells = table.add_row().cells
#     cells[0].paragraphs[0].add_run(str(count)).font.size = Cm(.34)
#     cells[1].paragraphs[0].add_run(str(str(k))).font.size = Cm(.34)
#     cells[2].paragraphs[0].add_run(str(str(v))).font.size = Cm(.34)
#     count += 1
        
# doc.add_page_break()
# doc.save('BookCount.docx')
# print("saved")
    


#     