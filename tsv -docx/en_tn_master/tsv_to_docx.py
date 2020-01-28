# -*- coding: utf-8 -*-
import docx 
import csv
import re
import os
import glob
from docx.shared import Cm 


def countword(wordcount):
    count = 0
    for words in wordcount:
        splitword = words.split()
        count = count + len(splitword)
    # print(count)
    return count

def wordcountdoc(word_count):
    # print(word_count)
    doc = docx.Document()
    table = doc.add_table(rows=0, cols=3)
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(3)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(5)
        section.right_margin = Cm(5)

    table.style = 'Table Grid'

    table.autofit = False 
    table.allow_autofit = False 

    table.columns[0].width = Cm(1.5) # Try to set column 0 width to 1.0
    table.columns[1].width = Cm(5)
    table.columns[2].width = Cm(4)



    heading = table.add_row().cells
    heading[0].paragraphs[0].add_run('No').bold = True
    heading[1].paragraphs[0].add_run('Book').bold = True
    heading[2].paragraphs[0].add_run('Count').bold = True
    count = 1
    for k, v in word_count.items():
        cells = table.add_row().cells
        cells[0].text = str(count)
        cells[1].text = k
        cells[2].text = str(v)
        count += 1
    doc.add_page_break()
    doc.save('BookCount.docx')
    print("saved")
        



files = glob.glob(os.getcwd() + "/en_tn-v14/en_tn/*.tsv")
# tcv_files = glob.glob('*.tsv')

word_count = {}
for fl in files:
  # filepath = re.sub(r"\.tsv","",fl)  bookname = fl.split("/")[-1].split(".")[0]
  print(bookname)
  filename = fl.split
  ("/")[-1]
  path = glob.glob(os.getcwd() + "/en_tn-v14/en_tn/" + filename)
 

  with open(path[0]) as tsvfile:
    reader = csv.reader(tsvfile, delimiter='\t')     
    doc = docx.Document()
    table = doc.add_table(rows=0, cols=1)

    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(1)
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)

    #----------Adding a heading -----------
    heading = doc.add_heading(bookname, level=1).alignment = 1
    lst = []
    wordcount = []
    table = doc.add_table(rows=0, cols=1)

    cells = table.add_row().cells
    cells[0].paragraphs[0].add_run("General Instructions:\n\n").bold = True
    # cells = table.add_row().cells
    cells[0].paragraphs[0].add_run("1. Translation of the content in ‘OccurenceNote’ should be strictly placed in the ‘Translation’ column only.\n").italic = True
    # cells = table.add_row().cells
    cells[0].paragraphs[0].add_run("2. ‘#’, ‘$’, ‘@’ are meta-tags that should be placed at the same positions in the translated text.\n").italic = True
    # cells = table.add_row().cells
    cells[0].paragraphs[0].add_run("3. Please do not modify any content in ‘Book’, ‘Chapter’, ‘Verse’, and ‘OccurenceNote’ columns.").italic = True
    table = doc.add_table(rows=0, cols=5)

    table.style = 'Table Grid'

    table.autofit = False 
    table.allow_autofit = False 

    table.columns[0].width = Cm(1.5) # Try to set column 0 width to 1.0
    table.columns[1].width = Cm(1.5)
    table.columns[2].width = Cm(1.5)
    table.columns[3].width = Cm(7.5)
    table.columns[4].width = Cm(8)

    heading_cells = table.add_row().cells
    heading_cells[0].paragraphs[0].add_run('Bk').bold = True
    heading_cells[1].paragraphs[0].add_run('Ch').bold = True
    heading_cells[2].paragraphs[0].add_run('Vs').bold = True
    heading_cells[3].paragraphs[0].add_run('OccurrenceNote').bold = True
    heading_cells[4].paragraphs[0].add_run('Translation').bold = True  
    for rows in reader:
    #   print(rows)
      lst.append(rows)
    for row in lst[1:]:
      book = row[0]
      chapter = row[1]
      verse = row[2]
      try:
        content = row[8]
      except:
        print("issue on row")
      # content = row[8]
      editcontent = re.sub(r"<br>","$",content)
      content1 = re.sub(r"(\[\[\w+\:[\/\w+\-]*\]\])","@",editcontent)

    # for counting words..
      rem_link = re.sub(r"(\(\w+\:\s\[\[\w+\:[\/\w+\-]*\]\]\))","",editcontent)
      rem_verse = re.sub(r"(\(\d+\:\d+\-\d+\))","", rem_link)
      rem_punctuation = re.sub(r'[^\w\s]','',rem_verse)
      rem_numbers = re.sub(r'\d+','',rem_punctuation)
      wordcount.append(rem_numbers)
    
      cells = table.add_row().cells
      cells[0].text = book
      cells[1].text = chapter
      cells[2].text = verse
      cells[3].text = content1 
      # cells[4].text = ''

    wordslength = countword(wordcount)
    word_count[bookname] = wordslength

    # print(word_count)
    
  doc.add_page_break()
  doc.save(bookname+'.docx')
  print("saved")

wordcountdoc(word_count)


# for italic and bold
# 
# p = doc.add_paragraph(style = 'List Bullet')
#     para1 = p.add_run('Translation of the content in ‘OccurenceNote’ should be strictly placed in the ‘Translation’ column only.')
#     para1.bold = True
#     para1.italic = True
#     doc.add_paragraph('‘#’, ‘$’, ‘@’ are meta-tags that should be placed at the same positions in the translated text.').style = 'List Bullet'
#     doc.add_paragraph('Please do not modify any content in ‘Book’, ‘Chapter’, ‘Verse’, and ‘OccurenceNote’ columns.').style = 'List Bullet'



