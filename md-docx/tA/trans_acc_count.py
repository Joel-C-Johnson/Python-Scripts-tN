# -*- coding: utf-8 -*-
import re
import os
import glob
import docx 
from docx.shared import Cm 


files = glob.glob(os.getcwd() + "/en_ta-v11/en_ta/process/*")
# print files
countw = {}
total_count = 0
for folder in files:
    wordcount = 0
    filename = folder.split('/')[-1]
    nextpath = glob.glob(folder+"/*.md")
    for fl in nextpath:
        f = open(fl, 'r')
        content = f.read()
        splitContent = content.split('\n')
        for lines in splitContent:
            filterContent = re.sub(r"(../.*?\.md)", "$", lines)
            rem_numbers = re.sub(r'\d+','',filterContent)
            rem_punctuation = re.sub(r'[^\w\s]','',rem_numbers)
            # output =  re.sub(r"\b[a-zA-Z]\b", "", rem_punctuation)
            finalContent = rem_punctuation.split()
            wordcount += len(finalContent)
    countw[filename] = wordcount
    total_count += wordcount
print(countw)
print(total_count)



doc = docx.Document()
heading = doc.add_heading("Word Count", level=2).alignment = 1
sections = doc.sections
for section in sections:
    section.top_margin = Cm(1)
    section.bottom_margin = Cm(1)
    section.left_margin = Cm(5)
    section.right_margin = Cm(5)


table = doc.add_table(rows=0, cols=3)
table.style = 'Table Grid'
table.autofit = False 
table.allow_autofit = False 
table.columns[0].width = Cm(1.15) 
table.columns[1].width = Cm(4)
table.columns[2].width = Cm(4)
heading1 = table.add_row().cells
heading1[0].paragraphs[0].add_run('No').bold = True
heading1[1].paragraphs[0].add_run('Book').bold = True
heading1[2].paragraphs[0].add_run('Count').bold = True

count = 1
for k, v in countw.items():
    cells = table.add_row().cells
    cells[0].paragraphs[0].add_run(str(count)).font.size = Cm(.34)
    cells[1].paragraphs[0].add_run(str(str(k))).font.size = Cm(.34)
    cells[2].paragraphs[0].add_run(str(str(v))).font.size = Cm(.34)
    count += 1
        
doc.add_page_break()
doc.save('twBookCount.docx')
print("saved")