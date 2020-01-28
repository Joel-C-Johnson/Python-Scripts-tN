# -*- coding: utf-8 -*-
import re
import os
import glob
import docx 
from docx.shared import Cm 


files = glob.glob(os.getcwd() + "/en_ta-v11/en_ta/process/*")
# print fil
for folder in files:
    nextpath = glob.glob(folder+"/*.md")
    # print(nextpath)
    doc = docx.Document()
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
    filesort =  []
    for fl in nextpath:
        filename = fl.split('/')[-3:]
        # filepath = "/".join(filename)   #checking/acceptable/sub-title.md
        if (filename[-1] == 'title.md'):
            filesort.insert(0, fl)
        elif (filename[-1] == 'sub-title.md'):
            filesort.insert(1, fl)
        else:
            filesort.insert(2, fl)
    # print(filesort)
    heading = doc.add_heading(filename[1], level=1).alignment = 1
    table = doc.add_table(rows=0, cols=1)
    table.style = 'Table Grid'
    cells = table.add_row().cells
    cells[0].paragraphs[0].add_run("General Instructions:\n").bold = True
    cells = table.add_row().cells
    cells[0].paragraphs[0].add_run("1. Translation of the content in 'English' should be strictly placed in the 'Translation' column only.").font.size = Cm(.36)    # cells = table.add_row().cells
    cells = table.add_row().cells
    cells[0].paragraphs[0].add_run("2. '#', '$', '@', '<u>' '</u>',  '*', '_' , [....], http://...... , etc. are meta-tags that should be placed at the same positions in the translated text.").font.size = Cm(.36)
    cells = table.add_row().cells
    cells[0].paragraphs[0].add_run("3. Please do not modify any content in 'English' columns.").font.size = Cm(.36)
    doc.add_paragraph('')
    # 
    for items in filesort:
        flname = items.split('/')[-3:]
        filepath = "/".join(flname)   #checking/acceptable/sub-title.md
# 
        f = open(items, 'r')
        content = f.read()
        splitContent = content.split('\n')
# 
        table = doc.add_table(rows=0, cols=1)
        cells = table.add_row().cells
        cells[0].paragraphs[0].add_run(filepath).bold = True
# 
        table = doc.add_table(rows=0, cols=3)
        table.style = 'Table Grid'
        table.autofit = False 
        table.allow_autofit = False 
        table.columns[0].width = Cm(1.15) 
        table.columns[1].width = Cm(8.35)
        table.columns[2].width = Cm(8.5)
# 
        heading1 = table.add_row().cells
        heading1[0].paragraphs[0].add_run('No').bold = True
        heading1[1].paragraphs[0].add_run('English').bold = True
        heading1[2].paragraphs[0].add_run('Translation').bold = True
        count = 1
        for lines in splitContent:
            filterContent = re.sub(r"(../.*?\.md)", "$", lines)
            cells = table.add_row().cells
            cells[0].paragraphs[0].add_run(str(count)).font.size = Cm(.34)
            cells[1].paragraphs[0].add_run(filterContent).font.size = Cm(.34)
            cells[2].text = ''
            count += 1
        doc.add_paragraph('')
        # doc.add_page_break()
            # 
        # 
    doc.save(filename[1]+'.docx')
    print("saved")
# 
# 
# 
# 
# 