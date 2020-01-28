# -*- coding: utf-8 -*-
import re
import os
import glob
import docx 
from docx.shared import Cm 


files = glob.glob(os.getcwd() + "/en_obs-tn-master/en_obs-tn/content/*")

# print(sorted(files))
countw = {}
total_count = 0

for folder in (sorted(files)):
    
    nextpath = glob.glob(folder+"/*.md")
    # print(sorted(nextpath))
    bookname = folder.split('/')[-1]
    wordcount = []
    for fl in sorted(nextpath):
        titles = fl.split('/')[-3:]
        titlePath = "/".join(titles)   #checking/acceptable/sub-title.md  
        print(titlePath)
        f = open(fl, 'r')
        content = f.read()
        splitContent = content.split('\n')
        for lines in splitContent:
            print lines
            rem_numbers = re.sub(r'\d+','',lines)
            rem_punctuation = re.sub(r'[^\w\s]','',rem_numbers)
            splitContent = rem_punctuation.split(" ")
            for words in splitContent:
                if words == '':
                    pass
                else:
                    wordcount.append(words)
    wordLength = len(wordcount)
    countw[bookname] = wordLength
    total_count += wordLength
print((countw))
print(total_count)      
            
#     # print("saved")



# doc = docx.Document()
# heading = doc.add_heading("Word Count", level=2).alignment = 1
# sections = doc.sections
# for section in sections:
#     section.top_margin = Cm(1)
#     section.bottom_margin = Cm(1)
#     section.left_margin = Cm(5)
#     section.right_margin = Cm(5)


# table = doc.add_table(rows=0, cols=3)
# table.style = 'Table Grid'
# table.autofit = False 
# table.allow_autofit = False 
# table.columns[0].width = Cm(1.15) 
# table.columns[1].width = Cm(4)
# table.columns[2].width = Cm(4)
# heading1 = table.add_row().cells
# heading1[0].paragraphs[0].add_run('No').bold = True
# heading1[1].paragraphs[0].add_run('Book').bold = True
# heading1[2].paragraphs[0].add_run('Count').bold = True

# count = 1
# for k, v in countw.items():
#     cells = table.add_row().cells
#     cells[0].paragraphs[0].add_run(str(count)).font.size = Cm(.34)
#     cells[1].paragraphs[0].add_run(str(str(k))).font.size = Cm(.34)
#     cells[2].paragraphs[0].add_run(str(str(v))).font.size = Cm(.34)
#     count += 1
        
# doc.add_page_break()
# doc.save('BookCount.docx')
# print("saved")