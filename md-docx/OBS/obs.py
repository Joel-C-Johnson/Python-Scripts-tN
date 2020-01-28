import re
import os
import glob
import docx 
from docx.shared import Cm 
# from docx.enum.text import WD_ALIGN_PARAGRAPH

files = glob.glob(os.getcwd() + "/Source/*.md")

def wordcountdoc(word_count, totalWordCount):
    # print(word_count)
    doc = docx.Document()
    table = doc.add_table(rows=0, cols=3)
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(5)
        section.right_margin = Cm(5)

    table.style = 'Table Grid'

    table.autofit = False 
    table.allow_autofit = False 

    table.columns[0].width = Cm(1.5) # Try to set column 0 width to 1.0
    table.columns[1].width = Cm(5)
    table.columns[2].width = Cm(4)



    heading = table.add_row().cells
    heading[0].paragraphs[0].add_run('No').bold = True
    heading[1].paragraphs[0].add_run('Book').bold = True
    heading[2].paragraphs[0].add_run('Count').bold = True
    count = 1
    for k, v in word_count.items():
        cells = table.add_row().cells
        cells[0].text = str(count)
        cells[1].text = k
        cells[2].text = str(v)
        count += 1
    cells = table.add_row().cells
    cells[1].paragraphs[0].add_run('Total Count').bold = True
    cells[2].paragraphs[0].add_run(str(sum(totalWordCount))).bold = True
    doc.add_page_break()
    doc.save('BookCount.docx')
    print("saved")


word_count = {}
totalWordCount = []
for fl in files:
    filename = fl.split("/")[-1]
    bookname = fl.split("/")[-1].split(".")[0]
    print(bookname)
    path = glob.glob(os.getcwd() + "/Source/" + filename)
    f = open(path[0], 'r')
    content = f.read()
    splitContent = content.split('\n')

    doc = docx.Document()
    heading = doc.add_heading(bookname, level=1).alignment = 1

    table = doc.add_table(rows=0, cols=1)

    cells = table.add_row().cells
    cells[0].paragraphs[0].add_run("General Instructions:\n").bold = True
    inst1 = cells[0].paragraphs[0].add_run("1. Translation of the content in 'English' should be strictly placed in the ‘Translation’ column only.\n").font.size = Cm(.33)    # cells = table.add_row().cells
    cells[0].paragraphs[0].add_run("2. ‘#’, ‘$’, ‘@’ are meta-tags that should be placed at the same positions in the translated text.\n").font.size = Cm(.33)
    cells[0].paragraphs[0].add_run("3. Please do not modify any content in 'English' columns.").font.size = Cm(.33)
    
    
    table = doc.add_table(rows=0, cols=3)
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    table.style = 'Table Grid'
    table.autofit = False 
    table.allow_autofit = False 
    table.columns[0].width = Cm(1.25) 
    table.columns[1].width = Cm(8.25)
    table.columns[2].width = Cm(8.5)

    heading1 = table.add_row().cells
    heading1[0].paragraphs[0].add_run('S.No').bold = True
    heading1[1].paragraphs[0].add_run('English').bold = True
    heading1[2].paragraphs[0].add_run('Translation').bold = True

    count = 1
    totalcount = 0
    for lines in splitContent:
        filterContent = re.sub(r"(https://.*?\.jpg)", "$", lines)
        
        #cout words
        rem_punctuation = re.sub(r'[^\w\s]','',filterContent)
        rem_numbers = re.sub(r'\d+','',rem_punctuation)
        # print(rem_numbers)
        coutWords = rem_numbers.split()
        totalcount = totalcount + len(coutWords)   
        # totalWordCount.append(totalcount) 

        cells = table.add_row().cells
        cells[0].paragraphs[0].add_run(str(count)).font.size = Cm(.33)
        cells[1].paragraphs[0].add_run(filterContent).font.size = Cm(.33)
        cells[2].text = ''
        count += 1
    word_count[filename] = totalcount
    totalWordCount.append(totalcount)
    doc.add_page_break()
    doc.save(bookname+'.docx')
    print("saved")

wordcountdoc(word_count, totalWordCount)

