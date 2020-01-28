import re
import os
import glob
import docx 
from docx.shared import Cm 
from docx import Document

Source_files = glob.glob(os.getcwd() + "/Source/obs/Content/*.md")

# print(Source_files)

Target_files = glob.glob(os.getcwd()+'/obs_intro/*.docx')
# print(Target_files)


for t_fl in Target_files:
    # t_bkname = t_fl.split('/')[-1]
    s_bkname = t_fl.split('/')[-1].split('.')[0]
    print(s_bkname)
    
    S_filePath = glob.glob(os.getcwd()+'/Source/obs/Content/' + s_bkname +'.md')
    fs = open(S_filePath[0], 'r')
    S_content = fs.read()
    findlink = re.findall(r'\!\[.*?\)',S_content)
    # print(findlink)
    
    fn = open(s_bkname +'.md', "w+")
    document = Document(t_fl)
    tables = document.tables
    source_list = []
    for table in tables:
        rows = table.rows
        for row in rows:
            try:
                tags = row.cells[1].text
                # print(tags)
                content = row.cells[2].text
                # print(content)
                source_list.append([tags, content])
            except:
                print("ksad")
    for t, v in source_list[1:]:
        if t == '':
            fn.write('\n')
            fn.write('\n')    
        elif t == '![OBS Image]($)':
            pop = findlink.pop(0)
            fn.write(pop)
        elif t and v =='':
            # pop = findlink.pop(0)
            fn.write(t)
        else:
            fn.write(v)
    fn.close()        
        









    # if tags:
                #     fn.write("hai")              
                # elif tags == '':
                #     fn.write("\n")
    # print(source_list)
    # fn.close()
    #         intro_list.append([i_tags, i_content])





    # S_filePath = glob.glob(os.getcwd()+'/Source/obs/Content/' + S_bkname)
    # fs = open(S_filePath[0], 'r')
    # S_content = fs.read()
    # findlink = re.findall(r'\!\[.*?\)',S_content)


    # T_filePath = glob.glob(os.getcwd()+'/Target/obs/OBS Hindi/files/')
    # print(T_filePath)
#     doc = docx.Document()
#     table = doc.add_table(rows=0, cols=3)
#     sections = doc.sections
#     for section in sections:
#         section.top_margin = Cm(1.5)
#         section.bottom_margin = Cm(1.5)
#         section.left_margin = Cm(5)
#         section.right_margin = Cm(5)

#     table.style = 'Table Grid'

#     table.autofit = False 
#     table.allow_autofit = False 

#     table.columns[0].width = Cm(1.5) # Try to set column 0 width to 1.0
#     table.columns[1].width = Cm(5)
#     table.columns[2].width = Cm(4)



#     heading = table.add_row().cells
#     heading[0].paragraphs[0].add_run('No').bold = True
#     heading[1].paragraphs[0].add_run('Book').bold = True
#     heading[2].paragraphs[0].add_run('Count').bold = True
#     count = 1
#     for k, v in word_count.items():
#         cells = table.add_row().cells
#         cells[0].text = str(count)
#         cells[1].text = k
#         cells[2].text = str(v)
#         count += 1
#     cells = table.add_row().cells
#     cells[1].paragraphs[0].add_run('Total Count').bold = True
#     cells[2].paragraphs[0].add_run(str(sum(totalWordCount))).bold = True
#     doc.add_page_break()
#     doc.save('BookCount.docx')
#     print("saved")


# word_count = {}
# totalWordCount = []
# for fl in files:
#     filename = fl.split("/")[-1]
#     bookname = fl.split("/")[-1].split(".")[0]
#     print(bookname)
#     path = glob.glob(os.getcwd() + "/FW__Bilingual_format_-_OBS/Content/" + filename)
#     f = open(path[0], 'r')
#     content = f.read()
#     splitContent = content.split('\n')

#     doc = docx.Document()
#     heading = doc.add_heading(bookname, level=1).alignment = 1

#     table = doc.add_table(rows=0, cols=1)

#     cells = table.add_row().cells
#     cells[0].paragraphs[0].add_run("General Instructions:\n").bold = True
#     inst1 = cells[0].paragraphs[0].add_run("1. Translation of the content in 'English' should be strictly placed in the ‘Translation’ column only.\n").font.size = Cm(.33)    # cells = table.add_row().cells
#     cells[0].paragraphs[0].add_run("2. ‘#’, ‘$’, ‘@’ are meta-tags that should be placed at the same positions in the translated text.\n").font.size = Cm(.33)
#     cells[0].paragraphs[0].add_run("3. Please do not modify any content in 'English' columns.").font.size = Cm(.33)
    
    
#     table = doc.add_table(rows=0, cols=3)
#     sections = doc.sections
#     for section in sections:
#         section.top_margin = Cm(2)
#         section.bottom_margin = Cm(2)
#         section.left_margin = Cm(2)
#         section.right_margin = Cm(2)

#     table.style = 'Table Grid'
#     table.autofit = False 
#     table.allow_autofit = False 
#     table.columns[0].width = Cm(1.25) 
#     table.columns[1].width = Cm(8.25)
#     table.columns[2].width = Cm(8.5)

#     heading1 = table.add_row().cells
#     heading1[0].paragraphs[0].add_run('S.No').bold = True
#     heading1[1].paragraphs[0].add_run('English').bold = True
#     heading1[2].paragraphs[0].add_run('Translation').bold = True

#     count = 1
#     totalcount = 0
#     for lines in splitContent:
#         filterContent = re.sub(r"(https://.*?\.jpg)", "$", lines)
        
#         #cout words
#         rem_punctuation = re.sub(r'[^\w\s]','',filterContent)
#         rem_numbers = re.sub(r'\d+','',rem_punctuation)
#         # print(rem_numbers)
#         coutWords = rem_numbers.split()
#         totalcount = totalcount + len(coutWords)   
#         # totalWordCount.append(totalcount) 

#         cells = table.add_row().cells
#         cells[0].paragraphs[0].add_run(str(count)).font.size = Cm(.33)
#         cells[1].paragraphs[0].add_run(filterContent).font.size = Cm(.33)
#         cells[2].text = ''
#         count += 1
#     word_count[filename] = totalcount
#     totalWordCount.append(totalcount)
#     doc.add_page_break()
#     doc.save(bookname+'.docx')
#     print("saved")

# wordcountdoc(word_count, totalWordCount)

