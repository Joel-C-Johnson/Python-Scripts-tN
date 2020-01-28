import re
import os
import glob
import docx
from docx.shared import Cm 
from docx.shared import Pt


files = glob.glob(os.getcwd() + "/IRV NT's/Urdu-NT/*.usfm")


for fl in files:
    f = open(fl, 'r')
    content = f.read()
    splitcontent = content.split('\n')
    # m = '\id JUD'
    # bk = re.match(r'(\\id )\w+',splitcontent).group()
    # print(splitcontent[0]) 
    doc = docx.Document()
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
    for lines in splitcontent:
        filename = re.match(r'(\\id )\w+',lines)
        bkname = re.match(r'(\\h )\w+',lines)
        chapter = re.match(r'(\\cl )', lines)
        chapter1 = re.match(r'(\\c )', lines)
        parag = re.match(r'(\\p)', lines)
        verse = re.match(r'(\\v )', lines)
        subheading = re.match(r'(\\s1 )', lines)
        subheading1 = re.match(r'(\\s )', lines)
        quotes = re.match(r'(\\q1 )', lines)
        quotes1 = re.match(r'(\\q2 )', lines)
        if bkname:
            bk = lines.split('\h')
            heading = doc.add_heading(bk[1], level=1).alignment = 1
            # print(bk[1])  
        elif chapter:
            # print(chapter)
            chptr = lines.split('\cl')
            heading = doc.add_heading(chptr[1], level=2).alignment = 0
            # print(chptr[1])
        elif chapter1:
            # print(chapter1)
            chptr = lines.split('\c')
            heading = doc.add_heading(chptr[1], level=2).alignment = 0
            # print(chptr[1])
        elif parag:
            paragraph = doc.add_paragraph('')
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = Pt(0)
            paragraph_format.space_after = Pt(0)
        elif verse:
            ver = lines.split('\\v')
            paragraph = doc.add_paragraph(ver[1])
            paragraph_format = paragraph.paragraph_format
            # p = doc.add_paragraph(ver[1])
            paragraph_format.space_before, paragraph_format.space_after
            paragraph_format.space_before = Pt(0)
            paragraph_format.space_after = Pt(6)
            # p.add_run(ver[1])
            # print(ver[1])
        elif subheading:
            sub = lines.split('\s1')
            heading = doc.add_heading(sub[1], level=3).alignment = 0
            # p = doc.add_paragraph().add_run(str('str')).bold = True
        elif subheading1:
            sub = lines.split('\s')
            heading = doc.add_heading(sub[1], level=3).alignment = 0
            # p = doc.add_paragraph().add_run(str('str')).bold = True
        elif filename:
            bkn = lines.split('\\id')
            fl = bkn[1].split(' ')[1]
            # print(fl) 
        elif quotes:
            qut = lines.split('\\q1')
            # print(qut)
            paragraph = doc.add_paragraph(qut[1])
            paragraph_format = paragraph.paragraph_format
            # p = doc.add_paragraph(ver[1])
            paragraph_format.space_before, paragraph_format.space_after
            paragraph_format.space_before = Pt(0)
            paragraph_format.space_after = Pt(6)

        elif quotes1:
            qut = lines.split('\\q2')
            # print(qut)
            paragraph = doc.add_paragraph(qut[1])
            paragraph_format = paragraph.paragraph_format
            # p = doc.add_paragraph(ver[1])
            paragraph_format.space_before, paragraph_format.space_after
            paragraph_format.space_before = Pt(0)
            paragraph_format.space_after = Pt(6)
        else:
            pass
    print(fl)  
    doc.save(fl+'.docx')
    print("saved")
            

    