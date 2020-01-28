from docx import Document
import glob
import os
import re
import docx 
from docx.shared import Cm 


files = glob.glob(os.getcwd() + "/ESV SN/03) ESV Cross-references/*.docx")
# print (files)
countw = {}
total_count = 0
for fl in files:    
    wordcount = 0  
    bknme = fl.split("/")[-1].split('.')[0].split(' ')
    bknme1 = " ".join(bknme)
    print (bknme1)
    document = Document(fl)
    tables = document.tables
    for table in tables:
        rows = table.rows
        for row in rows:    
            try:
                content = row.cells[4].text
                # print(content)
                if content == "English":
                    pass
                # rem_brackets = re.sub(r'\(.*?\)','',content)
                else:
                    tags = re.sub(r'\\\w+','',content)
                    rem_unclearText = re.sub(r' XML | Timeline | Timelinecharttimelinepngspan | v | ch | Na | eg | II | III | IV | V | VI | VII | VIII | IX | X | Dan | Tim | Ps | Kgs | OT | Heb | Deut | c | ad | Ex | Jer | Gen | Zech | Isa | Matt | Obad | Lam | Mic | Lev | Neh | Ezek | nd | Rom | Cor | Gal | km | Eph | ver | Phil | Acts | Rev | Job | Prov | UTF | ESV | Crossway | Num | Pet | Mal | Her | vv | bc | Nah | b | Chr | Sam | v | pdf | I | J | T | L | Hos | Ti | esv | SSB | Lk | Tm | Dt | Mt | Pt | Jn | Prv | Eccl | Dn | Jb | Ti | Zec | Tm | Mi | NT | F | DD | C | Gn | Rv | Jl | Mk | Col | NA ','', tags)
                    rem_numbers = re.sub(r'\d+','',rem_unclearText)
                    rem_punctuation = re.sub(r'[^\w\s]','',rem_numbers)
                    output =  re.sub(r"\b[a-zA-Z]\b", "", rem_punctuation)
                    splitContent = output.split(" ")
                    # print(splitContent)
                    counting = len(list(filter(None, splitContent)))
                    wordcount += counting
            except:
                pass
    countw[bknme1] = wordcount
    total_count += wordcount
print(countw)
print(total_count)


doc = docx.Document()
heading = doc.add_heading("Word Count", level=2).alignment = 1
sections = doc.sections
for section in sections:
    section.top_margin = Cm(1)
    section.bottom_margin = Cm(1)
    section.left_margin = Cm(5)
    section.right_margin = Cm(5)


table = doc.add_table(rows=0, cols=3)
table.style = 'Table Grid'
table.autofit = False 
table.allow_autofit = False 
table.columns[0].width = Cm(1.15) 
table.columns[1].width = Cm(4)
table.columns[2].width = Cm(4)
heading1 = table.add_row().cells
heading1[0].paragraphs[0].add_run('No').bold = True
heading1[1].paragraphs[0].add_run('Book').bold = True
heading1[2].paragraphs[0].add_run('Count').bold = True

count = 1
for k, v in countw.items():
    cells = table.add_row().cells
    cells[0].paragraphs[0].add_run(str(count)).font.size = Cm(.34)
    cells[1].paragraphs[0].add_run(str(str(k))).font.size = Cm(.34)
    cells[2].paragraphs[0].add_run(str(str(v))).font.size = Cm(.34)
    count += 1
        
doc.add_page_break()
doc.save('BookCount.docx')
print("saved")




