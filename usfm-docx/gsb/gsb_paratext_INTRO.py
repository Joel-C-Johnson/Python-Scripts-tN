# -*- coding: utf-8 -*-
import re
import os
import glob
import docx
from docx.shared import Cm 
from docx.shared import Pt


files = glob.glob(os.getcwd() + "/GSB_ParaText//*.SFM")

# print files
for fl in files:
    f = open(fl, 'r')
    content = f.read()
    splitcontent = content.split('\n')
    doc = docx.Document()
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
    for lines in splitcontent:
        bkname = re.match(r'(\\id )', lines)
        if bkname:
            bkn = lines.split(' ')
            print(bkn[1])
            print("\n" + "\n")
            heading = doc.add_heading(bkn[1], level=1).alignment = 1
            doc.add_paragraph()
            break
    table = doc.add_table(rows=0, cols=4)
    table.style = 'Table Grid'
    table.autofit = False 
    table.allow_autofit = False 
    table.columns[0].width = Cm(1.15) 
    table.columns[1].width = Cm(2)
    table.columns[2].width = Cm(7.5)
    table.columns[3].width = Cm(7.4)
    heading1 = table.add_row().cells
    heading1[0].paragraphs[0].add_run('No').bold = True
    heading1[1].paragraphs[0].add_run('Tags').bold = True
    heading1[2].paragraphs[0].add_run('English').bold = True
    heading1[3].paragraphs[0].add_run('Translation').bold = True
    count = 1
    for lines in splitcontent:
        # print (lines)
        chapter1 = re.match(r'(\\c )', lines)
        bkname = re.match(r'(\\id )', lines)
        ide = re.match(r'(\\ide )', lines)
        fig = re.search(r'(\\fig )', lines)
        if bkname:
            bkn = lines.split(' ')
            tags = bkn[0]
            content = " ".join(bkn[1:])
            cells = table.add_row().cells
            cells[0].paragraphs[0].add_run(str(count)).font.size = Cm(.34)
            cells[1].paragraphs[0].add_run(str(tags)).font.size = Cm(.34)
            cells[2].paragraphs[0].add_run((content)).font.size = Cm(.34)
            cells[3].paragraphs[0].add_run((content)).font.size = Cm(.34)
            count += 1
            heading = doc.add_heading(bkn[1], level=1).alignment = 1
            # print bkn[1]
        elif fig:
            fig = lines.split(' ')
            tags = fig[0]
            content = " ".join(fig[1:])
            # print (content)
            cells = table.add_row().cells
            cells[0].paragraphs[0].add_run(str(count)).font.size = Cm(.34)
            cells[1].paragraphs[0].add_run(str(tags)).font.size = Cm(.34)
            cells[2].paragraphs[0].add_run((content)).font.size = Cm(.34)
            cells[3].paragraphs[0].add_run((content)).font.size = Cm(.34)
            count += 1
        elif ide:
            ide = lines.split(' ')
            tags = ide[0]
            content = " ".join(ide[1:])
            cells = table.add_row().cells
            cells[0].paragraphs[0].add_run(str(count)).font.size = Cm(.34)
            cells[1].paragraphs[0].add_run(str(tags)).font.size = Cm(.34)
            cells[2].paragraphs[0].add_run((content)).font.size = Cm(.34)
            cells[3].paragraphs[0].add_run((content)).font.size = Cm(.34)
            count += 1
        elif chapter1:
            break
        else:
            intro = lines.split(' ')
            tags = intro[0]
            content = " ".join(intro[1:])
            # print content
            filterContent = re.sub(r"–", "-", content)
            filtercc = re.sub(r"—", "-", filterContent)
            filterc = re.sub(r"…", "...", filtercc)
            filte = re.sub(r"’", "'", filterc)
            filt = re.sub(r"”", "'", filte)  
            fil = re.sub(r"“", "'", filt)
            cells = table.add_row().cells
            cells[0].paragraphs[0].add_run(str(count)).font.size = Cm(.34)
            cells[1].paragraphs[0].add_run(str(tags)).font.size = Cm(.34)
            try:
                cells[2].paragraphs[0].add_run((fil)).font.size = Cm(.34)
            except:
                print(fil)
                print("\n")

            cells[3].text = ''
            count += 1
    doc.add_paragraph('')
    doc.save(bkn[1]+'.docx')
    print("saved")








            # filterContent = re.sub(r"–", "-", lines)
            # filtercc = re.sub(r"–", "-", filterContent)
            # filterc = re.sub(r"…", "...", filtercc)
            # filte = re.sub(r"’", "'", filterc)
            # filt = re.sub(r"”", "'", filte)
            # fil = re.sub(r"“", "'", filt)

        
            # 
            # print intro
#         filename = re.match(r'(\\id )\w+',lines)
#         bkname = re.match(r'(\\h )\w+',lines)
#         chapter = re.match(r'(\\cl )', lines)
#         chapter1 = re.match(r'(\\c )', lines)
#         parag = re.match(r'(\\p)', lines)
#         verse = re.match(r'(\\v )', lines)
#         subheading = re.match(r'(\\s1 )', lines)
#         subheading1 = re.match(r'(\\s )', lines)
#         quotes = re.match(r'(\\q1 )', lines)
#         quotes1 = re.match(r'(\\q2 )', lines)
#         if bkname:
#             bk = lines.split('\h')
#             heading = doc.add_heading(bk[1], level=1).alignment = 1
#             # print(bk[1])  
#         elif chapter:
#             # print(chapter)
#             chptr = lines.split('\cl')
#             heading = doc.add_heading(chptr[1], level=2).alignment = 0
#             # print(chptr[1])
#         elif chapter1:
#             # print(chapter1)
#             chptr = lines.split('\c')
#             heading = doc.add_heading(chptr[1], level=2).alignment = 0
#             # print(chptr[1])
#         elif parag:
#             paragraph = doc.add_paragraph('')
#             paragraph_format = paragraph.paragraph_format
#             paragraph_format.space_before = Pt(0)
#             paragraph_format.space_after = Pt(0)
#         elif verse:
#             ver = lines.split('\\v')
#             paragraph = doc.add_paragraph(ver[1])
#             paragraph_format = paragraph.paragraph_format
#             # p = doc.add_paragraph(ver[1])
#             paragraph_format.space_before, paragraph_format.space_after
#             paragraph_format.space_before = Pt(0)
#             paragraph_format.space_after = Pt(6)
#             # p.add_run(ver[1])
#             # print(ver[1])
#         elif subheading:
#             sub = lines.split('\s1')
#             heading = doc.add_heading(sub[1], level=3).alignment = 0
#             # p = doc.add_paragraph().add_run(str('str')).bold = True
#         elif subheading1:
#             sub = lines.split('\s')
#             heading = doc.add_heading(sub[1], level=3).alignment = 0
#             # p = doc.add_paragraph().add_run(str('str')).bold = True
#         elif filename:
#             bkn = lines.split('\\id')
#             fl = bkn[1].split(' ')[1]
#             # print(fl) 
#         elif quotes:
#             qut = lines.split('\\q1')
#             # print(qut)
#             paragraph = doc.add_paragraph(qut[1])
#             paragraph_format = paragraph.paragraph_format
#             # p = doc.add_paragraph(ver[1])
#             paragraph_format.space_before, paragraph_format.space_after
#             paragraph_format.space_before = Pt(0)
#             paragraph_format.space_after = Pt(6)

#         elif quotes1:
#             qut = lines.split('\\q2')
#             # print(qut)
#             paragraph = doc.add_paragraph(qut[1])
#             paragraph_format = paragraph.paragraph_format
#             # p = doc.add_paragraph(ver[1])
#             paragraph_format.space_before, paragraph_format.space_after
#             paragraph_format.space_before = Pt(0)
#             paragraph_format.space_after = Pt(6)
#         else:
#             pass
#     print(fl)  
#     doc.save(fl+'.docx')
#     print("saved")
            

    