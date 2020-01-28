# -*- coding: utf-8 -*-
import re
import os
import glob
import docx
from docx.shared import Cm 
from docx.shared import Pt


files = glob.glob(os.getcwd() + "/GSB_Source/GSB_ParaText/*.SFM")

# print files
for fl in files:
    f = open(fl, 'r')
    content = f.read()
    splitChapter = content.split('\n')
    doc = docx.Document()
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1.7)
        section.bottom_margin = Cm(1.7)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
    for lines in splitChapter:
        bkname = re.match(r'(\\id )', lines)
        if bkname:
            bkn = lines.split(' ')
            print(bkn[1])
            heading = doc.add_heading(bkn[1], level=1).alignment = 1
            doc.add_paragraph()
            break
    table = doc.add_table(rows=0, cols=1)
    table.style = 'Table Grid'
    cells = table.add_row().cells
    cells[0].paragraphs[0].add_run("General Instructions:\n").bold = True
    cells = table.add_row().cells
    cells[0].paragraphs[0].add_run("1. Translation of the content in 'English' should be strictly placed in the 'Translation' column only.").font.size = Cm(.36)    # cells = table.add_row().cells
    cells = table.add_row().cells
    cells[0].paragraphs[0].add_run("2. '\\x', '\\xo', '\\xt', '\\x*', .... etc. are usfm-tags that should be placed at the same positions in the translated text.").font.size = Cm(.36)
    cells = table.add_row().cells
    cells[0].paragraphs[0].add_run("3. Please do not modify any content in 'No', 'Ch', 'Vs', 'Tags', 'English' columns.").font.size = Cm(.36)
    doc.add_paragraph('')

    table = doc.add_table(rows=0, cols=6)
    table.style = 'Table Grid'
    table.autofit = False 
    table.allow_autofit = False 
    table.columns[0].width = Cm(1.1) 
    table.columns[1].width = Cm(1)
    table.columns[2].width = Cm(1)
    table.columns[3].width = Cm(3.5)
    table.columns[4].width = Cm(6)
    table.columns[5].width = Cm(6)
    heading1 = table.add_row().cells
    heading1[0].paragraphs[0].add_run('No').bold = True
    heading1[1].paragraphs[0].add_run('Ch').bold = True
    heading1[2].paragraphs[0].add_run('Vs').bold = True
    heading1[3].paragraphs[0].add_run('Tags').bold = True
    heading1[4].paragraphs[0].add_run('English').bold = True
    heading1[5].paragraphs[0].add_run('Translation').bold = True
    ch = 0
    
    count = 1
    for lines in splitChapter:
        chapter = re.match(r'(\\c )', lines)
        verse = re.match(r'(\\v )', lines)
        if chapter:
            ch = ch + 1
            vr = 0
        elif verse:
            vr = vr + 1
            fN = re.findall(r'(\\x \+.*?\\x\*)', lines) #(\\x +.*?\\xt)(.*)?(\\x\*)
            if fN:
                for i in fN:
                    if i == '':
                        pass
                    else:
                        tags = re.search(r'(\\x +.*?\\xt)(.*)?(\\x\*)', i)
                        cells = table.add_row().cells
                        cells[0].paragraphs[0].add_run(str(count)).font.size = Cm(.34)
                        cells[1].paragraphs[0].add_run(str(ch)).font.size = Cm(.34)
                        cells[2].paragraphs[0].add_run(str(vr)).font.size = Cm(.34)
                        if tags:
                            cells[3].paragraphs[0].add_run(str(tags.group(1) + " (text) " + tags.group(3))).font.size = Cm(.34)
                            cells[4].paragraphs[0].add_run(tags.group(2)).font.size = Cm(.34)
                            cells[5].text = ''
                        else:
                            print("tags Missing")
                        count += 1
            else:
                pass
    doc.add_paragraph('')
    doc.save(bkn[1]+'.docx')
    print("saved")


