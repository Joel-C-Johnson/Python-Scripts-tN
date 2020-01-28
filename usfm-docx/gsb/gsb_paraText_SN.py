# -*- coding: utf-8 -*-
import re
import os
import glob
import docx
from docx.shared import Cm 
from docx.shared import Pt


files = glob.glob(os.getcwd() + "/GSB_Source/GSB_ParaText/*.SFM")

# print files
for fl in files:
    f = open(fl, 'r')
    content = f.read()
    splitNewLine = content.split('\n')
    splitContent = content.split('\c ')
    join_content = ' '.join(splitContent[1:])
    source = join_content.split('\n')
    # print(source)
    doc = docx.Document()
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.6)
        section.left_margin = Cm(1.3)
        section.right_margin = Cm(1.3)
    for lines in splitNewLine:
        bkname = re.match(r'(\\id )', lines)
        if bkname:
            bkn = lines.split(' ')
            print(bkn[1])
            # print("\n" + "\n")
            heading = doc.add_heading(bkn[1], level=1).alignment = 1
            doc.add_paragraph()
            break
    table = doc.add_table(rows=0, cols=1)
    table.style = 'Table Grid'
    cells = table.add_row().cells
    cells[0].paragraphs[0].add_run("General Instructions:\n").bold = True
    cells = table.add_row().cells
    cells[0].paragraphs[0].add_run("1. Translation of the content in 'English' should be strictly placed in the 'Translation' column only.").font.size = Cm(.36)    # cells = table.add_row().cells
    cells = table.add_row().cells
    cells[0].paragraphs[0].add_run("2. '\\x', '\\xo', '\\xt', '\\x*', '\\fq*', '\\ef',  .... etc. are usfm-tags that should be placed at the same positions in the translated text.").font.size = Cm(.36)
    cells = table.add_row().cells
    cells[0].paragraphs[0].add_run("3. Please do not modify any content in 'No', 'Tags', 'Ch', 'Vs', 'Sub-Tags', 'English' columns.").font.size = Cm(.36)
    doc.add_paragraph('')

    table = doc.add_table(rows=0, cols=7)
    table.style = 'Table Grid'
    table.autofit = False 
    table.allow_autofit = False 
    table.columns[0].width = Cm(1.1) 
    table.columns[1].width = Cm(1.4)
    table.columns[2].width = Cm(1)
    table.columns[3].width = Cm(1)
    table.columns[4].width = Cm(2.5)
    table.columns[5].width = Cm(6)
    table.columns[6].width = Cm(6)
    heading1 = table.add_row().cells
    heading1[0].paragraphs[0].add_run('No').bold = True
    heading1[1].paragraphs[0].add_run('Tags').bold = True
    heading1[2].paragraphs[0].add_run('Ch').bold = True
    heading1[3].paragraphs[0].add_run('Vs').bold = True
    heading1[4].paragraphs[0].add_run('Sub-Tags').bold = True
    heading1[5].paragraphs[0].add_run('English').bold = True
    heading1[6].paragraphs[0].add_run('Translation').bold = True
    
    
    ch = 0
    count = 1
    for lines in source:
        Strip_chapter = lines.strip(' ')
        chapter = Strip_chapter.isnumeric()
        verse = re.match(r'(\\v )', lines)
        if chapter == True:
            ch = ch + 1
            print(ch)
            vr = 0
            cells = table.add_row().cells
            cells[0].paragraphs[0].add_run(str(count)).font.size = Cm(.34)
            cells[1].paragraphs[0].add_run("\c").font.size = Cm(.34)
            cells[2].paragraphs[0].add_run(str(ch)).font.size = Cm(.34)
            cells[3].paragraphs[0].add_run(str(vr)).font.size = Cm(.34)
            cells[4].paragraphs[0].add_run("NA").font.size = Cm(.34)
            cells[5].paragraphs[0].add_run("NA").font.size = Cm(.34)
            cells[6].paragraphs[0].add_run("NA").font.size = Cm(.34)
            count += 1
        elif verse:
            # print(verse.group())
            vr = vr + 1
            sub_tags = re.findall(r'(\\ef - .*?\\ef\*)', lines)
            # print(sub_tags)
            if sub_tags:
                for i in sub_tags:
                    if i == '':
                        pass
                    else:
                        tags = re.search(r'(\\ef -.*?\\fr.*?\\\w+)(.*)?(\\ef\*)', i)
                        # print(tags.group())
                        tag1 = re.search(r'(\\ef -.*?\\fig)(.*)?', i)   #\\ef - .*?\\fig
                        if tags:
                            cells = table.add_row().cells
                            cells[0].paragraphs[0].add_run(str(count)).font.size = Cm(.34)
                            cells[1].paragraphs[0].add_run(str(verse.group())).font.size = Cm(.34)
                            cells[2].paragraphs[0].add_run(str(ch)).font.size = Cm(.34)
                            cells[3].paragraphs[0].add_run(str(vr)).font.size = Cm(.34)
                            cells[4].paragraphs[0].add_run(str(tags.group(1) + " (text) " + tags.group(3))).font.size = Cm(.34)
                            cells[5].paragraphs[0].add_run(tags.group(2)).font.size = Cm(.34)
                            cells[6].text = ''
                            count += 1
                        elif tag1:
                            cells = table.add_row().cells
                            cells[0].paragraphs[0].add_run(str(count)).font.size = Cm(.34)
                            cells[1].paragraphs[0].add_run(str(verse.group())).font.size = Cm(.34)
                            cells[2].paragraphs[0].add_run(str(ch)).font.size = Cm(.34)
                            cells[3].paragraphs[0].add_run(str(vr)).font.size = Cm(.34)
                            cells[4].paragraphs[0].add_run(tag1.group(0)).font.size = Cm(.34)
                            cells[5].paragraphs[0].add_run(tag1.group(0)).font.size = Cm(.34)
                            cells[6].paragraphs[0].add_run(tag1.group(0)).font.size = Cm(.34)
                            count += 1
                        else:
                            print("tags ssss")
                            # print(sub_tags)
            else:
                split_verse = re.search(r'(\\v )(\d+)(.*)', lines)
                # print(split_verse.group(0)) 
                cells = table.add_row().cells
                cells[0].paragraphs[0].add_run(str(count)).font.size = Cm(.34)
                cells[1].paragraphs[0].add_run(str(split_verse.group(1))).font.size = Cm(.34)
                cells[2].paragraphs[0].add_run(str(ch)).font.size = Cm(.34)
                cells[3].paragraphs[0].add_run(str(vr)).font.size = Cm(.34)
                cells[4].paragraphs[0].add_run("NA").font.size = Cm(.34)
                cells[5].paragraphs[0].add_run("NA").font.size = Cm(.34)
                cells[6].paragraphs[0].add_run("NA").font.size = Cm(.34)
                count += 1
        
        elif lines:
            splitLine = re.split(' ', lines)
            # print(splitLine[0])
            # print(' '.join(splitLine[1:]))
            # print(len(splitLine))
            if len(splitLine) == 1:
                cells = table.add_row().cells
                cells[0].paragraphs[0].add_run(str(count)).font.size = Cm(.34)
                cells[1].paragraphs[0].add_run(splitLine[0]).font.size = Cm(.34)
                cells[2].paragraphs[0].add_run(str(ch)).font.size = Cm(.34)
                cells[3].paragraphs[0].add_run(str(vr)).font.size = Cm(.34)
                cells[4].paragraphs[0].add_run("NA").font.size = Cm(.34)
                cells[5].paragraphs[0].add_run("NA").font.size = Cm(.34)
                cells[6].paragraphs[0].add_run("NA").font.size = Cm(.34)
                count += 1
            else:
                cells = table.add_row().cells
                cells[0].paragraphs[0].add_run(str(count)).font.size = Cm(.34)
                cells[1].paragraphs[0].add_run(splitLine[0]).font.size = Cm(.34)
                cells[2].paragraphs[0].add_run(str(ch)).font.size = Cm(.34)
                cells[3].paragraphs[0].add_run(str(vr)).font.size = Cm(.34)
                cells[4].paragraphs[0].add_run("NA").font.size = Cm(.34)
                cells[5].paragraphs[0].add_run(' '.join(splitLine[1:])).font.size = Cm(.34)
                cells[6].text = ''
                count += 1
                        
        else:
            print("tags Missing")
            # print(lines)
        # print("Program Running")
                        
    doc.add_paragraph('')
    doc.save(bkn[1]+'.docx')
    print("saved")
