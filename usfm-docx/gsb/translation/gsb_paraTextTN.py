# -*- coding: utf-8 -*-
import re
import os
import glob
import docx
from docx.shared import Cm
from docx.shared import Pt
import openpyxl
from docx import Document

tn_files = glob.glob(os.getcwd() + "/tn/*.docx")

intro = "Introductions.docx"
sn = "Study Notes etc.docx"
cr = "Cross Reference.docx"

# ----------------------------INTRO-----------------------------#

in_path = glob.glob(os.getcwd() + "/tn/" + str(intro))
in_document = Document(in_path[0])
in_tables = in_document.tables
intro_list = []
for i_table in in_tables:
    i_rows = i_table.rows
    for i_row in i_rows:
        i_tags = i_row.cells[1].text
        i_content = i_row.cells[3].text
        intro_list.append([i_tags, i_content])

fn = open("3JN.usfm", "w+")
for k, v in intro_list[1:]:
    fn.write(k + " " + v + "\n")
fn.close()

#---------------------------CR---------------------------------#

cr_list = []
cr_path = glob.glob(os.getcwd() + "/tn/" + str(cr))
cr_document = Document(cr_path[0])
cr_tables = cr_document.tables
for cr_table in cr_tables:
    cr_rows = cr_table.rows
    for cr_row in cr_rows:
        try:
            cr_tags = cr_row.cells[3].text
            cr_content = cr_row.cells[5].text
            # print(content)
            tags_replace = re.sub(r'\(text\)', cr_content, cr_tags)
            cr_list.append(tags_replace)
        except:
            pass
# print(cr_list)

#-----------------------------SN----------------------------#

sn_list = []
sn_path = glob.glob(os.getcwd() + "/tn/" + str(sn))
sr_document = Document(sn_path[0])
sr_tables = sr_document.tables
for sr_table in sr_tables:
    sr_rows = sr_table.rows
    for sr_row in sr_rows:
        try:
            sr_tags = sr_row.cells[4].text
            sr_content = sr_row.cells[6].text
            # print(sr_content)
            sr_slash = re.sub(r'\\', '$$', sr_content)
            # print("sr",sr_slash)
            sr_tags_replace = re.sub(r'\(text\)|NA', sr_slash, sr_tags)
            sn_list.append(sr_tags_replace)
        except:
            pass
sn_list.pop(0)
# print("snList",sn_list)
# target_path = glob.glob(os.getcwd() +"/" +"3JN.usfm")
# f= open(target_path[0],"a+")
# for i in sn_list:
#     f.write(i + "\n")
# f.close()


# --------------------------SOURCE-----------------------------#

files = glob.glob(os.getcwd() + "/source/*.SFM")

# print(files)
target_path = glob.glob(os.getcwd() + "/" + "3JN.usfm")
ft = open(target_path[0], "a+")
for fl in files:
    f = open(fl, 'r')
    content = f.read()
    splitNewLine = content.split('\n')
    splitContent = content.split('\c ')
    join_content = ' '.join(splitContent[1:])
    s2 = re.sub(r'\\x.*?\\x\*', " @@ ", join_content)
    for item in cr_list[1:]:
        s2 = s2.replace(r' @@ ', item, 1)
    source = s2.split('\n')
    # print(source)
    ch = 0
    count = 1
    for lines in source:
        Strip_chapter = lines.strip(' ')
        chapter = Strip_chapter.isnumeric()
        verse = re.match(r'(\\v )', lines)
        esb = re.match(r'\\esb ', lines)
        esbe = re.match(r'\\esbe', lines)
        if chapter == True:
            ch = ch + 1
            ft.write("\c" + " " + str(ch) + "\n")
            sn_list.pop(0)
        elif verse:
            sub_tags = re.findall(r'(\\ef - .*?\\ef\*)', lines)
            # find_cr = re.findall(r'\\x.*?\\x\*', lines)
            # print(find_cr)
            if sub_tags:
                # print(sub_tags)
                for i in sub_tags:
                    # print(i)
                    # [1,2]
                    if i == '':
                        pass
                    else:
                        tags = re.search(r'(\\ef - .*?\\fr.*?\\\w+)(.*)?(\\ef\*)', lines)
                        tag1 = re.search(r'(\\ef -.*?\\fig)(.*)?', i)
                        if tags:
                            pop = sn_list.pop(0)
                            if pop:
                                lines = lines.replace( i, pop, 1)
                        if tag1:
                            pop = sn_list.pop(0)
                            if pop:
                                lines = lines.replace( i, pop, 1)
                                # c2 = c1.replace('$$','\\')
                                # ft.write(c2 + "\n")
                c2 = lines.replace('$$','\\')
                ft.write(c2 + "\n")
            else:
                sn_pop = sn_list.pop(0)
                c1 = lines.replace( lines, sn_pop)
                c2 = c1.replace('$$','\\')
                ft.write(c2 + "\n")

        elif esb:
            ft.write("\esb" + "\n")
            sn_list.pop(0)
        elif esbe:
            ft.write("\esbe" + "\n")
            sn_list.pop(0)
        elif lines:
            splitLine = re.split(' ', lines)
            if len(splitLine) == 1:
                ft.write(splitLine[0] + "\n")
                sn_list.pop(0)
            else:
                replace_tags = sn_list.pop(0)
                re_line = replace_tags.replace('$$', '\\')
                ft.write(splitLine[0] + " " + re_line + "\n")
        else:
            print("No tags")
    ft.close()
