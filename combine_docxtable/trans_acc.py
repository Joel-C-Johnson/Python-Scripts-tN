# -*- coding: utf-8 -*-
import re
import os
import glob
import docx 
from docx.shared import Cm 
from docx import Document
from collections import OrderedDict

files = glob.glob(os.getcwd() + "/tA_translate/*")
for folder in files:
    nextpath = glob.glob(folder+'/*.docx')
    filesort =  []
    dic = OrderedDict()
    for fl in nextpath:
        filename = fl.split('/')[-3:]
        # filepath = "/".join(filename)   #checking/acceptable/sub-title.md
        title_search = re.match(r'title',filename[-1])
        sub_title = re.match(r'sub',filename[-1])
        md_search = re.match(r'01',filename[-1])
        if title_search:
            filesort.insert(0, fl)
        elif sub_title:
            filesort.insert(1, fl)
        else:
            filesort.insert(2, fl)
        # if (filename[-1] == 'title.md_kan.docx'):
        #     filesort.insert(0, fl)
        # elif (filename[-1] == 'sub-title.md_kan.docx'):
        #     filesort.insert(1, fl)
        # else:
        #     filesort.insert(2, fl)
    print(filename[1])
    # print(filesort)
    for items in list(filesort):
        flname = items.split('/')[-3:]
        filepath1 = "/".join(flname)   #checking/acceptable/sub-title.md
        document = Document(items)
        tables = document.tables
        contentlist = []
        for table in tables:
            rows = table.rows
            for row in rows:
                try:
                    id = row.cells[0].text
                    # print(id)
                    content = row.cells[1].text
                    translation = row.cells[2].text
                    if(id==content and id==translation):
                        pass
                    else:
                        contentlist.append([id,content,translation])
                except:
                    print("ksad")
        dic[filepath1] = contentlist
  
    doc = Document()
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(1)
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)
    heading = doc.add_heading(filename[1], level=1).alignment = 1
    doc.add_paragraph()

        
    for k, v in dic.items():
        table = doc.add_table(rows=0, cols=1)
        table.style = 'Table Grid'
        cells = table.add_row().cells
        cells[0].paragraphs[0].add_run(k).bold = True
        len_v = 0
        for i in v:
            if len_v == 0:
                table = doc.add_table(rows=0, cols=3)
                table.style = 'Table Grid'
                table.autofit = False 
                table.allow_autofit = False 
                table.columns[0].width = Cm(1.25) # Try to set column 0 width to 1.0
                table.columns[1].width = Cm(9)
                table.columns[2].width = Cm(9.5)

                heading_cells = table.add_row().cells
                heading_cells[0].paragraphs[0].add_run(i[0]).bold = True
                heading_cells[1].paragraphs[0].add_run(i[1]).bold = True
                heading_cells[2].paragraphs[0].add_run(i[2]).bold = True

                len_v = 1
            else:
                table.style = 'Table Grid'
                table.autofit = False 
                table.allow_autofit = False 
                table.columns[0].width = Cm(1.25) # Try to set column 0 width to 1.0
                table.columns[1].width = Cm(9)
                table.columns[2].width = Cm(9.5)

                cells = table.add_row().cells
                cells[0].text = i[0].split()[0]
                cells[1].text = i[1]
                cells[2].text = i[2]

        doc.add_paragraph('')
    doc.add_page_break()
    doc.save(filename[1]+'.docx')
    # print("saved")